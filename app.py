from flask import Flask, request, jsonify, render_template, session, send_file
from flask_pymongo import PyMongo
from bson.objectid import ObjectId
from flask_cors import CORS
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import logging
import os

import os


app = Flask(__name__)
CORS(app)  # Enable CORS for cross-origin requests

# MongoDB configuration
app.config["MONGO_URI"] = "mongodb://localhost:27017/feedback_db"
mongo = PyMongo(app)

# Define collections
feedback_collection = mongo.db.feedbacks
subjects_collection = mongo.db.subjects
questions_collection = mongo.db.questions


@app.route("/")
def home():
    return render_template("user.html")


@app.route('/admin-login')
def admin_login():
    return render_template("admin.html")

# Rating text-to-number mapping
rating_map = {"Bad": 1, "Average": 2, "Good": 3, "Very Good": 4}

@app.route("/submit-feedback", methods=["POST"])
def submit_feedback():
    data = request.json
    register_number = data.get("register_number")
    semester = data.get("semester")
    branch = data.get("branch")
    subject = data.get("subject")
    form_type = data.get("form")
    year = data.get("year")
    ratings = data.get("ratings")

    # Check if feedback already exists for the same register number, branch, semester, subject, and year
    existing_feedback = feedback_collection.find_one({
        "register_number": register_number,
        "semester": semester,
        "branch": branch,
        "subject": subject,
        "year": year
    })

    if existing_feedback:
        return jsonify({"error": "Feedback already submitted for this subject and year!"}), 400  # Bad request

    # Insert new feedback
    feedback_collection.insert_one({
        "register_number": register_number,
        "semester": semester,
        "branch": branch,
        "subject": subject,
        "form": form_type,
        "year": year,
        "ratings": ratings
    })

    return jsonify({"message": "Feedback submitted successfully!"}), 200


@app.route('/fetch-feedback', methods=['GET'])
def fetch_feedback():
    branch = request.args.get("branch")
    semester = request.args.get("semester")
    subject = request.args.get("subject")
    form = request.args.get("form")
    year = request.args.get("year")  # Added year filter

    if not (branch and semester and subject and form and year):
        return jsonify({"error": "Missing parameters"}), 400

    # Fetch feedback from MongoDB for the selected year
    feedback_data = list(feedback_collection.find(
        {"branch": branch, "semester": semester, "subject": subject, "form": form, "year": year},
        {"_id": 0, "ratings": 1}
    ))

    if not feedback_data:
        return jsonify({"message": "No feedback available"}), 404

    # Fetch questions
    question_doc = questions_collection.find_one({"_id": "questions_collection"})
    if not question_doc:
        return jsonify({"error": "No questions found!"}), 404

    questions_key = f"{form}_questions"
    if questions_key not in question_doc:
        return jsonify({"error": f"No questions found for form type: {form}"}), 404

    questions = question_doc[questions_key]

    # Compute feedback summary
    result = []
    for i, question in enumerate(questions):
        rating_counts = {1: 0, 2: 0, 3: 0, 4: 0}

        for feedback in feedback_data:
            if i < len(feedback["ratings"]):
                rating = feedback["ratings"][i]
                if isinstance(rating, str) and rating in rating_map:
                    rating = rating_map[rating]
                elif isinstance(rating, (int, float)):
                    rating = int(rating)
                else:
                    continue  # Skip invalid ratings

                rating_counts[rating] += 1

        total_responses = sum(rating_counts.values())
        avg_rating = sum(rating * count for rating, count in rating_counts.items()) / total_responses if total_responses else 0

        result.append({
            "Question": question,
            "1 (Bad)": rating_counts[1],
            "2 (Average)": rating_counts[2],
            "3 (Good)": rating_counts[3],
            "4 (Very Good)": rating_counts[4],
            "AVG": round(avg_rating, 2)
        })

    return jsonify({
        "branch": branch,
        "semester": semester,
        "subject": subject,
        "form": form,
        "year": year,
        "feedback_data": result
    }), 200


    
    # Add the new question to the existing array of questions in the collection
    questions_collection.update_one(
        {"_id": ("questions_collection")},  # Use the appropriate document ID
        {"$push": {"questions": new_question}},
        upsert=True  # Create a new document if one does not exist
    )

    return jsonify({"message": "Question added successfully!"}), 201




@app.route('/get-questions', methods=['GET'])
def get_questions():
    question_doc = questions_collection.find_one({"_id": "questions_collection"})
    if not question_doc:
        return jsonify({"error": "No questions found!"}), 404

    return jsonify({
        "mid_questions": question_doc.get("mid_questions", []),
        "end_questions": question_doc.get("end_questions", [])
    }), 200



@app.route('/add-subject', methods=['POST'])
def add_subject():
    try:
        data = request.json
        branch = data.get("branch")
        semester = data.get("semester")
        new_subject = data.get("subject")

        if not (branch and semester and new_subject):
            return jsonify({"error": "Branch, Semester, and Subject are required!"}), 400

        # Append the new subject to the array of the specified semester
        subjects_collection.update_one(
            {"branch": branch},
            {"$push": {f"semesters.{semester}": new_subject}},
            upsert=True
        )

        return jsonify({"message": "Subject added successfully!"}), 201
    except Exception as e:
        print(f"Error: {e}")  # Debugging line
        return jsonify({"error": str(e)}), 500



@app.route('/get-subjects/<branch>/<semester>', methods=['GET'])
def get_subjects(branch, semester):
    subjects_data = subjects_collection.find_one({"branch": branch})
    
    if not subjects_data or semester not in subjects_data["semesters"]:
        return jsonify({"error": "No subjects found for this semester."}), 404

    subjects_list = subjects_data["semesters"][semester]
    return jsonify(subjects_list), 200

@app.route('/get-all-subjects', methods=['GET'])
def get_all_subjects():
    # Fetch all subjects from the database
    all_subjects = subjects_collection.find()
    subjects_data = []
    
    for subject_doc in all_subjects:
        branch = subject_doc.get("branch", "Unknown Branch")
        semesters = subject_doc.get("semesters", {})
        
        formatted_subjects = []
        for semester, subjects in semesters.items():
            formatted_subjects.append({
                "semester": semester,
                "subjects": subjects
            })
        
        subjects_data.append({
            "branch": branch,
            "semesters": formatted_subjects
        })
    
    return jsonify(subjects_data), 200
    
@app.route('/get-branches', methods=['GET'])
def get_branches():
    branches = subjects_collection.distinct("branch")
    return jsonify(branches), 200

@app.route('/delete-subject', methods=['DELETE'])
def delete_subject():
    data = request.json
    branch = data.get('branch')
    semester = str(data.get('semester'))  # Ensure semester is a string
    subject = data.get('subject')

    if not branch or not semester or not subject:
        return jsonify({"success": False, "message": "Missing required fields!"}), 400

    # Find the document for the specified branch
    subject_doc = subjects_collection.find_one({"branch": branch})
    
    if not subject_doc:
        return jsonify({"success": False, "message": "Branch not found!"}), 404

    # Check if the semester exists in the document
    if semester not in subject_doc.get("semesters", {}):
        return jsonify({"success": False, "message": "Semester not found!"}), 404

    # Check if the subject exists in the semester list
    if subject not in subject_doc["semesters"][semester]:
        return jsonify({"success": False, "message": "Subject not found!"}), 404

    # Remove the subject from the semester list
    subjects_collection.update_one(
        {"branch": branch},
        {"$pull": {f"semesters.{semester}": subject}}
    )

    return jsonify({"success": True, "message": "Subject deleted successfully!"}), 200

@app.route('/get-all-questions', methods=['GET'])
def get_all_questions():
    # Fetch all questions from the database
    question_doc = questions_collection.find_one({"_id": "questions_collection"})
    if not question_doc or "questions" not in question_doc:
        return jsonify({"error": "No questions found!"}), 404

    return jsonify(question_doc["questions"]), 200

@app.route('/download-feedback', methods=['GET'])
def download_feedback():
    branch = request.args.get("branch")
    semester = request.args.get("semester")
    subject = request.args.get("subject")
    form = request.args.get("form")  # 'mid' or 'end'
    year = request.args.get("year")  # Added year parameter

    # Validate all parameters
    if not all([branch, semester, subject, form, year]):
        return jsonify({"error": "Missing parameters"}), 400

    # Fetch feedback from MongoDB for the selected year
    feedback_data = list(feedback_collection.find(
        {"branch": branch, "semester": semester, "subject": subject, "form": form, "year": year},
        {"_id": 0, "ratings": 1}
    ))

    if not feedback_data:
        return jsonify({"message": "No feedback available"}), 404

    # Fetch questions from MongoDB
    question_doc = questions_collection.find_one({"_id": "questions_collection"})
    if not question_doc:
        return jsonify({"error": "No questions found!"}), 404

    questions_key = f"{form}_questions"
    if questions_key not in question_doc:
        return jsonify({"error": f"No questions found for form type: {form}"}), 404

    questions = question_doc[questions_key]

    # Create a Word document
    file_stream = create_feedback_template(subject, semester, branch, year, feedback_data, questions, form)
    return send_file(
        file_stream,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"feedback_{branch}{semester}{subject}_{form}_{year}.docx"
    )
def create_feedback_template(subject, semester, branch, year, feedback_data, questions, formtype):
    doc = Document()

    # Add logo to the header (first page only)
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Define the logo path
    # logo_path = "images/logo.jpg"  # Ensure this path is correct

    # Add logo (if file exists)
    # if logo_path:
    #     try:
    #         run = header_paragraph.add_run()
    #         run.add_picture(logo_path, width=Inches(1.5))  # Adjust width as needed
    #     except FileNotFoundError:
    #         header_paragraph.add_run("Logo Not Found").bold = True
    # else:
    #     header_paragraph.add_run("Logo Path Not Provided").bold = True

    # Add header text below the logo (only on the first page)
    header_text = header_paragraph.add_run("\nSTUDENT FEEDBACK FORM")
    header_text.bold = True

    # Add institution details (bold)
    institution_paragraph = doc.add_paragraph()
    institution_paragraph.add_run("Institution Name: East West Polytechnic").bold = True
    institution_paragraph.add_run("\nInstitution Code: 499").bold = True

    # Add a details table (no borders)
    table_details = doc.add_table(rows=1, cols=2)
    table_details.autofit = False
    
    # Left column: Branch, Semester, Subject (bold)
    left_cell = table_details.rows[0].cells[0]
    left_cell.text = f"Branch: {branch}\nSemester: {semester}\nSubject: {subject}"
    for paragraph in left_cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

    # Right column: Form Type, Year, Teacher’s Name (without extra space)
    right_cell = table_details.rows[0].cells[1]
    right_paragraph = right_cell.paragraphs[0]
    right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_paragraph.add_run(f"Form Type: {formtype}\nYear: {year}\nTeacher’s Name:").bold = True

    # Remove borders for the details table
    for row in table_details.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))  # Remove borders

    # Ensure details appear only on the first page
    for section in doc.sections[1:]:
        section.header.is_linked_to_previous = False
        section.header.paragraphs.clear()

    # Add feedback table
    table = doc.add_table(rows=1, cols=7)  # Extra column for Serial Number and Average
    table.style = "Table Grid"

    # Set column widths
    feedback_col_widths = [Inches(0.5), Inches(3.5), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0)]
    for i, width in enumerate(feedback_col_widths):
        table.columns[i].width = width

    # Add table headers (bold and center-aligned)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "S.No"
    hdr_cells[1].text = "Question"
    hdr_cells[2].text = "1\n(Bad)"
    hdr_cells[3].text = "2\n(Average)"
    hdr_cells[4].text = "3\n(Good)"
    hdr_cells[5].text = "4\n(Very Good)"
    hdr_cells[6].text = "Average"

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                if cell in hdr_cells[2:6]:  # Reduce font size for rating headers
                    run.font.size = Pt(8)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add questions and ratings to the table
    total_avg_sum = 0
    rating_map = {"Bad": 1, "Average": 2, "Good": 3, "Very Good": 4}  # Map ratings

    for i, question in enumerate(questions):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)  # Serial Number
        row_cells[1].text = question

        rating_counts = {1: 0, 2: 0, 3: 0, 4: 0}
        total_responses = 0
        total_rating = 0

        for feedback in feedback_data:
            if i < len(feedback["ratings"]):
                raw_rating = feedback["ratings"][i]

                if isinstance(raw_rating, str) and raw_rating in rating_map:
                    rating = rating_map[raw_rating]
                elif isinstance(raw_rating, (int, float)):
                    rating = int(raw_rating)
                else:
                    continue  # Skip invalid ratings

                rating_counts[rating] += 1
                total_responses += 1
                total_rating += rating

        row_cells[2].text = str(rating_counts[1])
        row_cells[3].text = str(rating_counts[2])
        row_cells[4].text = str(rating_counts[3])
        row_cells[5].text = str(rating_counts[4])

        avg_rating = total_rating / total_responses if total_responses > 0 else 0
        row_cells[6].text = f"{avg_rating:.2f}"
        total_avg_sum += avg_rating

        for cell in row_cells[2:]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add the total average box
    total_avg = total_avg_sum / len(questions) if len(questions) > 0 else 0
    total_avg_paragraph = doc.add_paragraph()
    total_avg_run = total_avg_paragraph.add_run(f"Total Average: {total_avg:.2f}")
    total_avg_run.bold = True
    total_avg_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save document to a BytesIO object
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream
    
if __name__ == "__main__":
    app.run(debug=True)



